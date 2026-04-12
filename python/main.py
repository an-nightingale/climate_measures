import threading
import time
import pandas as pd
from fastapi import FastAPI, HTTPException
from pydantic import BaseModel
from llama_index.llms.nvidia import NVIDIA
from llama_index.core.llms import ChatMessage
from llama_index.embeddings.huggingface import HuggingFaceEmbedding
from llama_index.core import Settings, VectorStoreIndex, Document, StorageContext
from llama_index.vector_stores.postgres import PGVectorStore
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
import nest_asyncio
from dotenv import load_dotenv
import os
import psycopg2
import openai
from docx import Document as DocxDocument
from docx.shared import Pt
from docx.enum.table import WD_TABLE_ALIGNMENT
from io import BytesIO
import re
from bs4 import BeautifulSoup
import openpyxl.styles
import json
from typing import Optional
from functools import lru_cache
import numpy as np
from rapidfuzz import fuzz

load_dotenv()
nest_asyncio.apply()

# === Глобальное хранилище одобренных мероприятий (в памяти) ===
APPROVED_MEASURES = []
_rebuild_lock = threading.Lock()
TABLE_NAME = "climate_embeddings"

yandex_client = openai.OpenAI(
    api_key=os.getenv("YANDEX_CLOUD_API_KEY"),
    base_url="https://ai.api.cloud.yandex.net/v1",
    project=os.getenv("YANDEX_CLOUD_FOLDER")
)

embed_model = HuggingFaceEmbedding(model_name='intfloat/multilingual-e5-large-instruct')
Settings.embed_model = embed_model
Settings.llm = NVIDIA(
    model="openai/gpt-oss-120b",
    api_key=os.getenv("NVIDIA_API_KEY"),
    max_tokens=8000,
)

# === Параметры подключения к PostgreSQL ===
DB_PARAMS = {
    "database": os.getenv("DB_DATABASE"),
    "host": os.getenv("DB_HOST"),
    "password": os.getenv("DB_PASSWORD"),
    "port": os.getenv("DB_PORT"),
    "user": os.getenv("DB_USERNAME"),
    "embed_dim": 1024,
}
PSYCOPG_DB_PARAMS = {
    "database": os.getenv("DB_DATABASE"),
    "host": os.getenv("DB_HOST"),
    "password": os.getenv("DB_PASSWORD"),
    "port": os.getenv("DB_PORT"),
    "user": os.getenv("DB_USERNAME"),
}

# === Промпты ===
RAG_SYSTEM_PROMPT = """
Ты — эксперт по адаптации к изменениям климата.
У тебя есть база знаний с кейсами и нормативными документами.
Пользователь вводит запрос, связанный с климатическим риском в регионе или отрасли.
Твоя задача — на основе информации из базы знаний предложить 2–3 релевантных адаптационных мероприятия,
которые помогут снизить климатический риск, о котором спрашивает пользователь.
### Требования к ответу:
1. Представь результат **в виде Markdown-таблицы** с колонками:
   - Наименование мероприятий
   - Митигационный эффект
   - Адаптационный эффект
   - Актуальность для региона (указать с учётом контекста запроса). Если регион не указан, считай, что задается вопрос по Тюменской области
   - Ответственная организация (из региона)
2. Если источник данных, на которых ты основываешь ответ, известен (это URL и краткое название кейса),
   добавь их **ниже таблицы** в виде списка ссылок:
   `**Опорные источники:** [1] Наименовавание мероприятий - URL, [2] Наименовавание мероприятий - URL`
3. Пиши кратко, по существу, с акцентом на реальные, практические меры.
Пример формата ответа:
| Наименование мероприятий | Митигационный эффект | Адаптационный эффект | Актуальность для Тобольского района | Ответственная организация |
|---------------------------|----------------------|----------------------|------------------------------------|----------------------------|
| Развитие городского электротранспорта | снижение эмиссии | повышение устойчивости транспортной инфраструктуры | актуально | городские власти |
| Перевод транспорта на газомоторное топливо | снижение эмиссии | рациональное использование ресурсов | реализуется частично | транспортные организации |
**Опорные источники:** [1] Наименовавание мероприятия 1 - https://example.com/case_1    [2] Наименовавание мероприятия 2 - https://example.com/case_2   
Приводи те источники, которые используешь для формирования таблицы непосредственно. Источники только из полученного контекста базы знаний, при этом приведи источник соответственно для каждой рекомендуемой тобой меры (сколько строк таблицы, столько и источников). 
URL в источниках приводи строго такое же, как указано в базе знаний. Наименование мероприятий в источниках бери из базы знаний. В таблице наименования мероприятий формируй без упоминания источников (других кейсов).
Ответственную организацию в таблице указывай актуальную для региона, который пользователь указал в запросе
В описаниях мероприятий в таблице не упоминай географические названия (рек, городов и тд) и не упоминай напрямую ничего про кейсы
"""

DIALOG_SYSTEM_PROMPT = """
Ты — экспертный ассистент по вопросам климатических рисков для Тюменской области. Отвечай на вопросы пользователя максимально точно и полезно. Ищи актуальную информацию. После поиска дай пользователю краткий, структурированный ответ, выдели ключевые факты. Если использовал внешние источники - в конце напиши ссылки на страницы, с которых взята информация.
"""

CLASSIFIER_SYSTEM_PROMPT = """
Ты — классификатор запросов. Твоя задача — определить тип запроса пользователя.
Доступные типы:
- "rag" — запросы, требующие поиска в базе знаний, напрямую связанные с адаптационными мероприятиями для конкретных регионов. Могут быть либо прямыми запросами на составление адаптационных мероприятий, либо описанием существующего климатического риска с просьбой составить адаптационные мероприятия или без просьбы.
- "dialog" — прочие вопросы, не требующие поиска в базе знаний, в том числе общие вопросы про климатические риски.
- "statistics" — запросы на получение статистики из базы данных по районам, периодам, показателям и единицам измерения.

Примеры классификации:
Запрос: "Какие меры по адаптации к засухе для Ялуторовского района?"
Тип: rag

Запрос: "Составь таблицу мер по адаптации к аномальной жаре для Ишима"
Тип: rag

Запрос: "Пожарная опасность в Тобольском районе. Какие адаптационные мероприятия к ней?"
Тип: rag

Запрос: "Опасность паводков в Заводоуковске"
Тип: rag

Запрос: "Какие есть климатические риски в Тюменской области"
Тип: dialogue

Запрос: "Кто отвечает за адаптацию в Исетском районе?"
Тип: dialogue

Запрос: "Проведи анализ климатических рисков для Аббатского района, связанных с засухой"
Тип: dialog

Запрос: "Расскажи про адаптацию в Норвегии"
Тип: dialog

Запрос: "Оборот розничной торговли (без субъектов малого предпринимательства) в Абатском районе за 2022 г. в млн. руб."
Тип: statistics

Запрос: "Вывезено за год твердых коммунальных отходов в Викуловском районе за 2023 г. в тыс. тонн"
Тип: statistics

Запрос: "Покажи численность населения в Армизонском районе на 1 января 2025 года"
Тип: statistics

ВАЖНО: Если во входных данных есть "История диалога", учитывай её. 
Если пользователь ссылается на предыдущий ответ, который относится к типу rag (например, "добавь ещё меры к таблице выше"), это тип "rag".
Если пользователь ссылается на предыдущий статистический ответ, просит продолжить, уточнить, сравнить или вывести похожий показатель — это тип "statistics".
Если пользователь просто благодарит или задает вопрос на новую тему  — "dialog".
Отвечай ТОЛЬКО одним словом: "rag", "dialog" или "statistics". Не добавляй никаких пояснений.
"""
STATISTICS_SQL_SYSTEM_PROMPT = """
Ты — эксперт по PostgreSQL и аналитике статистических данных.
Сформируй ОДИН безопасный SQL SELECT-запрос к базе статистики.

Правила:
1. Возвращай только SQL, без пояснений и без Markdown.
2. Разрешён только SELECT.
3. Нельзя использовать INSERT, UPDATE, DELETE, DROP, ALTER, TRUNCATE, CREATE.
4. Используй только таблицы и поля из описания схемы.
5. Используй значения районов, периодов, секций и индикаторов только из переданного контекста.
6. Район нужно выбирать ВСЕГДА явно через territory.name. Если в запросе пользователя указан район, обязательно добавляй фильтр по territory.name.
7. Не фильтруй по unit.name, даже если пользователь указал желаемую единицу измерения. Единицы измерения нужно вернуть в результирующих строках, а не использовать как фильтр.
8. Для периода не ориентируйся на period.name как основной способ фильтрации. Основной способ — period.start_date и period.end_date.
9. Если пользователь просит значение "за год", считай конечной датой начало следующего года (01.01). Захватывай год полностью, не обрывай раньше времени, если есть показатель за начало последующего года. Отбирай период по диапазону дат:
   - period.start_date >= DATE '2025-01-01'
   - period.end_date <= DATE '2026-01-01'
10. Если пользователь просит "на 1 января 2025 года" или другую дату, то ищи по start_date и end_date этой даты:
   - period.start_date = DATE '2025-01-01'
   - period.end_date = DATE '2025-01-01'
11. Всегда возвращай полную информацию по найденным строкам:
   - territory.name AS "Территория"
   - indicator.name AS "Показатель"
   - section.name AS "Секция"
   - industry.name AS "Категория"
   - unit.name AS "Единица измерения"
   - period_type.name AS "Тип периода"
   - period.name AS "Период"
   - period.start_date AS "Дата начала"
   - period.end_date AS "Дата окончания"
   - statistic.value AS "Значение"
12. Обычно нужны JOIN:
   - statistic -> territory, indicator, period
   - indicator -> section, unit
   - section -> industry
   - period -> period_type
13. Если пользователь просит одно значение, всё равно возвращай полную строку со всеми полями.
14. Если пользователь просит несколько районов, сравнение районов, сумму по районам, разницу между районами, сумму за несколько периодов или разницу между периодами — допускаются простые арифметические операции SQL:
   - SUM(...)
   - разность двух агрегатов
   - CASE WHEN
   - GROUP BY
   - CTE / WITH
15. Если пользователь просит разницу между годами, сравни значения за нужные периоды и верни вычисленный результат, но также по возможности приложи исходные строки.
16. Если пользователь просит сумму показателей двух районов, суммируй statistic.value по нужным районам.
17. Для поиска показателя ориентируйся прежде всего на indicator.name.
18. Возвращай понятные русские псевдонимы столбцов.
19. Если пользователь просит значение "за год", считай конечной датой начало следующего года (01.01). Захватывай год полностью, не обрывай раньше времени, если есть показатель за начало последующего года. Отбирай период по диапазону дат:
   - period.start_date >= DATE '2025-01-01'
   - period.end_date <= DATE '2026-01-01'
Возвращай только SQL.
"""

STATISTICS_ANSWER_SYSTEM_PROMPT = """
Ты — ассистент по статистике муниципальных районов.
Тебе переданы:
1. Исходный запрос пользователя.
2. Краткое описание таблиц и полей БД.
3. Результат SQL-запроса в виде строк "название столбца: значение".

Сформируй краткий ответ на русском языке. Не давай пояснений или таблиц, напиши четкий полный ответ на поставленный вопрос - включи все пункты из самого вопроса и непосредственный ответ. Выделяй важные числа жирным шрифтом, только сами числа.

Правила:
- Не выдумывай значения, используй только переданные результаты.
- Если строк несколько, оформи результат списком или Markdown-таблицей.
- Если результат пустой, честно скажи, что по заданным параметрам данные не найдены, и предложи уточнить показатель, период или район.
- Не показывай пользователю внутренние технические детали SQL.
- Если пользователь просил единицу измерения, а в результатах единица другая, выполни перевод единиц в искомые пользователем, если это возможно сделать однозначно.
- При переводе единиц явно укажи итоговое пересчитанное значение и целевую единицу измерения.
- Если корректно перевести единицы невозможно или неоднозначно, честно скажи об этом и покажи исходные значения как есть.
- Если запрос пользователя был на разницу, сумму или другое простое вычисление, объясни результат коротко и понятно.
"""

STATISTICS_SCHEMA_DESCRIPTION = """
Таблицы БД статистики:
- industry(industry_id, name) — отрасль/категория секции.
- territory_type(territory_type_id, name) — тип территории.
- unit(unit_id, name) — единица измерения показателя.
- period_type(period_type_id, name) — тип периода, например "период" или "дата".
- territory(territory_id, parent_territory_id, territory_type_id, name) — территория, например район.
- section(section_id, industry_id, name) — раздел статистики.
- indicator(indicator_id, section_id, unit_id, name) — показатель.
- period(period_id, period_type_id, name, start_date, end_date) — период или дата.
- statistic(statistic_id, territory_id, indicator_id, period_id, value) — числовое значение.

Основные связи:
- statistic.territory_id -> territory.territory_id
- statistic.indicator_id -> indicator.indicator_id
- statistic.period_id -> period.period_id
- indicator.section_id -> section.section_id
- indicator.unit_id -> unit.unit_id
- section.industry_id -> industry.industry_id
- period.period_type_id -> period_type.period_type_id
- territory.territory_type_id -> territory_type.territory_type_id
"""


# === Утилиты для работы с Excel ===
def read_excel_as_documents(file_path: str):
    """
    Создаёт документы для векторного индекса:
    - Текст для эмбеддинга: только первые 4 столбца
    - Метаданные: остальные столбцы + служебная информация
    """
    df = pd.read_excel(file_path)
    docs = []

    # Столбцы для векторизации (по ним будет поиск)
    embed_columns = ['Проблема', 'Наименование мероприятий', 'Митигационный эффект', 'Адаптационный эффект']
    # Столбцы для метаданных (добавляются в контекст после поиска)
    meta_columns = ['Наименование района', 'Агроклиматические условия района', 'Ответственная организация', 'Источник']

    for i, row in df.iterrows():
        # === Текст для эмбеддинга (только первые 4 колонки) ===
        embed_text_parts = []
        for col in embed_columns:
            if col in df.columns and pd.notna(row[col]):
                embed_text_parts.append(f"{col}: {row[col]}")
        embed_text = "\n".join(embed_text_parts)

        # === Метаданные (остальные колонки + служебные) ===
        metadata = {
            "source": os.path.basename(file_path),
            "row_index": i,
            "file_type": "excel"
        }
        for col in meta_columns:
            if col in df.columns and pd.notna(row[col]):
                metadata[f"meta_{col}"] = str(row[col])

        docs.append(Document(
            text=embed_text,
            metadata=metadata
        ))
    return docs


def process_excel_files(data_path: str):
    all_docs = []
    for file_name in os.listdir(data_path):
        if file_name.endswith(('.xlsx', '.xls')):
            file_path = os.path.join(data_path, file_name)
            print(f"Обработка Excel файла: {file_name}")
            try:
                excel_docs = read_excel_as_documents(file_path)
                all_docs.extend(excel_docs)
                print(f"  - Получено {len(excel_docs)} документов")
            except Exception as e:
                print(f"  - Ошибка обработки {file_name}: {e}")
    return all_docs


# === Преобразование APPROVED_MEASURES в документы ===
def get_approved_documents():
    """
    Преобразует APPROVED_MEASURES в документы.
    Для консистентности: первые 4 поля - для поиска, остальные - в метаданные.
    """
    docs = []
    for i, item in enumerate(APPROVED_MEASURES):
        # Текст для эмбеддинга (первые 4 поля)
        embed_text = "\n".join([
            f"Проблема: {item.get('source_question', '')}",
            f"Наименование мероприятий: {item.get('name', '')}",
            f"Митигационный эффект: {item.get('mitigation', '')}",
            f"Адаптационный эффект: {item.get('adaptation', '')}"
        ])

        # Метаданные (остальные поля)
        metadata = {
            "source": "user_approved_in_memory",
            "row_index": i,
            "file_type": "user_approved",
            "meta_Наименование района": item.get('relevance', ''),
            "meta_Ответственная организация": item.get('responsible', ''),
            "meta_Источник": ""  # можно добавить поле для URL при одобрении
        }

        docs.append(Document(
            text=embed_text,
            metadata=metadata
        ))
    return docs


# === Фоновый ребилд индекса (БЕЗ ПРОСТОЯ) ===
def background_rebuild_index():
    if _rebuild_lock.locked():
        print("🔁 Ребилд уже запущен")
        return

    with _rebuild_lock:
        print("🔄 Запуск фонового ребилда...")

        # Имя новой таблицы
        temp_table = f"climate_embeddings_new_{int(time.time())}"
        active_table = TABLE_NAME
        backup_table = f"climate_embeddings_old_{int(time.time())}"

        try:
            # 1. Создаём новую таблицу и наполняем данными
            temp_store = PGVectorStore.from_params(
                table_name=temp_table,
                hnsw_kwargs={
                    "hnsw_m": 16,
                    "hnsw_ef_construction": 64,
                    "hnsw_ef_search": 40,
                    "hnsw_dist_method": "vector_cosine_ops",
                },
                **DB_PARAMS
            )

            # 2. Собираем документы
            all_docs = []
            all_docs.extend(process_excel_files("./data"))
            all_docs.extend(get_approved_documents())

            if not all_docs:
                raise ValueError("Нет документов для индексации")

            # 3. Создаём индекс в новой таблице
            storage_context = StorageContext.from_defaults(vector_store=temp_store)
            VectorStoreIndex(
                nodes=all_docs,
                storage_context=storage_context,
                show_progress=True
            )
            print(f"  → Новая таблица {temp_table} готова")

            # 4. АТОМАРНО ПЕРЕКЛЮЧАЕМ ТАБЛИЦЫ
            conn = psycopg2.connect(dbname=os.getenv("DB_DATABASE"), user=os.getenv("DB_USERNAME"),
                                    password=os.getenv("DB_PASSWORD"), host=os.getenv("DB_HOST"),
                                    port=os.getenv("DB_PORT"))
            cur = conn.cursor()
            cur.execute(f"ALTER TABLE IF EXISTS data_{active_table} RENAME TO {backup_table};")
            cur.execute(f"ALTER TABLE {temp_table} RENAME TO data_{active_table};")
            conn.commit()
            cur.close()
            conn.close()

            print(f"✅ Ребилд завершён. Старая таблица: {backup_table}")

        except Exception as e:
            print("❌ Ошибка ребилда:", e)


# === Загрузка индекса (для RAG) ===
def get_vector_store():
    return PGVectorStore.from_params(
        table_name=TABLE_NAME,
        hnsw_kwargs={
            "hnsw_m": 16,
            "hnsw_ef_construction": 64,
            "hnsw_ef_search": 40,
            "hnsw_dist_method": "vector_cosine_ops",
        },
        **DB_PARAMS
    )


def load_vector_index():
    try:
        vector_store = get_vector_store()
        index = VectorStoreIndex.from_vector_store(vector_store=vector_store)
        return index
    except Exception as e:
        print(f"Ошибка загрузки индекса: {e}")
        return None


# === Модифицированная функция диалогового агента с Яндекс GPT ===
def generate_dialog_response(user_question: str, conversation_history: str = None) -> str:
    try:
        history_instruction = ""
        if conversation_history:
            history_instruction = f"\n\nИстория диалога:\n{conversation_history}\n\nУчитывай предыдущие сообщения."

        full_prompt = f"{DIALOG_SYSTEM_PROMPT}{history_instruction}\nЗапрос пользователя: {user_question}"
        response = yandex_client.responses.create(
            model=f"gpt://{os.getenv('YANDEX_CLOUD_FOLDER')}/{os.getenv('YANDEX_CLOUD_MODEL')}",
            input=full_prompt,
            tools=[
                {
                    "type": "web_search",
                    "filters": {
                        "allowed_domains": [],
                        "user_location": {
                            "region": "225",
                        }
                    },
                    "search_context_size": "medium",
                }
            ],
            temperature=0.5,
            max_output_tokens=2000
        )
        if hasattr(response, 'output_text') and response.output_text:
            return response.output_text
        else:
            return f"Не удалось получить ответ от Яндекс GPT. Пожалуйста, попробуйте еще раз."

    except Exception as e:
        print(f"Ошибка диалогового агента Яндекс GPT: {e}")
        return f"Ошибка при обработке запроса: {str(e)}"


# === Статистический агент ===
def get_db_connection():
    return psycopg2.connect(**PSYCOPG_DB_PARAMS)


@lru_cache(maxsize=1)
def get_statistics_metadata(_cache_key: int):
    conn = get_db_connection()
    try:
        indicators_df = pd.read_sql_query(
            """
            SELECT i.name AS indicator_name,
                   s.name AS section_name,
                   u.name AS unit_name,
                   COALESCE(ind.name, '') AS industry_name
            FROM indicator i
            JOIN section s ON s.section_id = i.section_id
            JOIN unit u ON u.unit_id = i.unit_id
            LEFT JOIN industry ind ON ind.industry_id = s.industry_id
            ORDER BY s.name, i.name
            """,
            conn,
        )
        territories_df = pd.read_sql_query(
            """
            SELECT t.name AS territory_name,
                   COALESCE(tt.name, '') AS territory_type
            FROM territory t
            LEFT JOIN territory_type tt ON tt.territory_type_id = t.territory_type_id
            ORDER BY t.name
            """,
            conn,
        )
        periods_df = pd.read_sql_query(
            """
            SELECT p.name AS period_name,
                   COALESCE(pt.name, '') AS period_type,
                   p.start_date,
                   p.end_date
            FROM period p
            LEFT JOIN period_type pt ON pt.period_type_id = p.period_type_id
            ORDER BY p.name
            """,
            conn,
        )
        units_df = pd.read_sql_query("SELECT name AS unit_name FROM unit ORDER BY name", conn)
        sections_df = pd.read_sql_query(
            """
            SELECT s.name AS section_name,
                   COALESCE(ind.name, '') AS industry_name
            FROM section s
            LEFT JOIN industry ind ON ind.industry_id = s.industry_id
            ORDER BY s.name
            """,
            conn,
        )
        return {
            "indicators": indicators_df,
            "territories": territories_df,
            "periods": periods_df,
            "units": units_df,
            "sections": sections_df,
        }
    finally:
        conn.close()


def invalidate_statistics_metadata_cache():
    get_statistics_metadata.cache_clear()


def statistics_cache_key() -> int:
    return int(time.time() // 300)


def get_top_similar_territories(user_question: str, top_k: int = 2, min_score: float = 70.0) -> pd.DataFrame:
    metadata = get_statistics_metadata(statistics_cache_key())
    territories_df = metadata["territories"].copy()

    def score_row(row) -> float:
        return float(fuzz.partial_ratio(user_question.lower(), str(row["territory_name"]).lower()))

    territories_df["similarity"] = territories_df.apply(score_row, axis=1)

    top_df = (
        territories_df
        .sort_values("similarity", ascending=False)
        .head(top_k)
        .reset_index(drop=True)
    )

    top_df = top_df[top_df["similarity"] >= min_score].reset_index(drop=True)
    return top_df


def get_indicators_for_territories(territory_names: list[str]) -> pd.DataFrame:
    if not territory_names:
        metadata = get_statistics_metadata(statistics_cache_key())
        return metadata["indicators"].copy()

    conn = get_db_connection()
    try:
        sql = """
        SELECT DISTINCT
               i.name AS indicator_name,
               s.name AS section_name,
               u.name AS unit_name,
               COALESCE(ind.name, '') AS industry_name,
               t.name AS territory_name
        FROM statistic st
        JOIN territory t ON t.territory_id = st.territory_id
        JOIN indicator i ON i.indicator_id = st.indicator_id
        JOIN section s ON s.section_id = i.section_id
        JOIN unit u ON u.unit_id = i.unit_id
        LEFT JOIN industry ind ON ind.industry_id = s.industry_id
        WHERE t.name = ANY(%s)
        ORDER BY i.name
        """
        return pd.read_sql_query(sql, conn, params=(territory_names,))
    finally:
        conn.close()


def cosine_similarity_matrix(query_vec: np.ndarray, matrix: np.ndarray) -> np.ndarray:
    query_norm = np.linalg.norm(query_vec)
    matrix_norms = np.linalg.norm(matrix, axis=1)

    if query_norm == 0:
        return np.zeros(len(matrix), dtype=float)

    safe_denominator = np.where(matrix_norms == 0, 1e-12, matrix_norms)
    sims = matrix @ query_vec / (safe_denominator * query_norm)
    return sims


@lru_cache(maxsize=1)
def get_indicator_embeddings(_cache_key: int):
    metadata = get_statistics_metadata(_cache_key)
    indicators_df = metadata["indicators"].copy()

    indicator_texts = []
    for _, row in indicators_df.iterrows():
        indicator_texts.append(
            f"Показатель: {row['indicator_name']}. "
            f"Секция: {row['section_name']}. "
            f"Единица: {row['unit_name']}. "
            f"Категория: {row['industry_name']}."
        )

    embeddings = Settings.embed_model.get_text_embedding_batch(indicator_texts)
    embeddings_matrix = np.array(embeddings, dtype=np.float32)

    return indicators_df, embeddings_matrix


def get_top_similar_indicators(user_question: str, top_k: int = 30) -> tuple[pd.DataFrame, pd.DataFrame]:
    matched_territories_df = get_top_similar_territories(user_question)
    matched_territory_names = matched_territories_df["territory_name"].dropna().astype(str).tolist()

    indicators_df = get_indicators_for_territories(matched_territory_names).copy()

    if indicators_df.empty:
        metadata = get_statistics_metadata(statistics_cache_key())
        indicators_df = metadata["indicators"].copy()

    def score_row(row) -> float:
        candidate = " | ".join([
            str(row["indicator_name"]),
            str(row["section_name"]),
            str(row["unit_name"]),
            str(row["industry_name"]),
        ])
        return float(fuzz.token_sort_ratio(user_question.lower(), candidate.lower()))

    indicators_df["similarity"] = indicators_df.apply(score_row, axis=1)

    top_df = (
        indicators_df
        .sort_values("similarity", ascending=False)
        .head(top_k)
        .reset_index(drop=True)
    )

    return matched_territories_df, top_df


def format_statistics_context(user_question: str) -> str:
    metadata = get_statistics_metadata(statistics_cache_key())

    matched_territories_df, top_indicators_df = get_top_similar_indicators(user_question, top_k=30)
    print("\n--- TOP SIMILAR TERRITORIES ---")
    for _, row in matched_territories_df.iterrows():
        print(f"{row['similarity']:.2f} | {row['territory_name']}")

    print("\n--- TOP SIMILAR INDICATORS ---")
    for _, row in top_indicators_df.iterrows():
        print(f"{row['similarity']:.4f} | {row['indicator_name']} | {row['section_name']} | {row['unit_name']}")
    print("--- END TOP SIMILAR INDICATORS ---\n")

    indicators_lines = []
    for _, row in top_indicators_df.iterrows():
        indicators_lines.append(
            f"- {row['indicator_name']} | секция: {row['section_name']} | единица: {row['unit_name']} | категория: {row['industry_name']} | similarity: {row['similarity']:.4f}"
        )

    territories_lines = []
    for _, row in metadata["territories"].head(500).iterrows():
        territories_lines.append(f"- {row['territory_name']} | тип: {row['territory_type']}")

    periods_lines = []
    for _, row in metadata["periods"].head(500).iterrows():
        periods_lines.append(
            f"- {row['period_name']} | тип: {row['period_type']} | start_date: {row['start_date']} | end_date: {row['end_date']}"
        )

    units_lines = [f"- {u}" for u in metadata["units"]["unit_name"].dropna().astype(str).tolist()]
    sections_lines = [
        f"- {row['section_name']} | категория: {row['industry_name']}"
        for _, row in metadata["sections"].head(500).iterrows()
    ]

    return (
            f"{STATISTICS_SCHEMA_DESCRIPTION}\n\n"
            f"Доступные территории:\n" + "\n".join(territories_lines) + "\n\n"
                                                                        f"Доступные периоды:\n" + "\n".join(
        periods_lines) + "\n\n"
                         f"Доступные единицы измерения:\n" + "\n".join(units_lines) + "\n\n"
                                                                                      f"Доступные секции:\n" + "\n".join(
        sections_lines) + "\n\n"
                          f"Доступные индикаторы:\n" + "\n".join(indicators_lines)
    )


def extract_sql_from_llm_response(text: str) -> str:
    return text.strip().replace("```sql", "").replace("```", "").strip()


def validate_statistics_sql(sql: str) -> str:
    sql_clean = sql.strip().rstrip(";")
    sql_lower = sql_clean.lower()

    if not (sql_lower.startswith("select") or sql_lower.startswith("with")):
        raise ValueError("Статистический агент может выполнять только SELECT-запросы")

    padded = f" {sql_lower} "
    for token in [" insert ", " update ", " delete ", " drop ", " alter ", " truncate ", " create ", " grant ",
                  " revoke "]:
        if token in padded:
            raise ValueError("Обнаружен запрещённый SQL-оператор")

    if ";" in sql_clean:
        raise ValueError("Разрешён только один SQL-запрос")

    return sql_clean


def generate_statistics_sql(user_question: str, conversation_history: str = None) -> str:
    context = format_statistics_context(user_question)
    history_block = f"\n\nИстория диалога:\n{conversation_history}" if conversation_history else ""
    messages = [
        ChatMessage(role="system", content=STATISTICS_SQL_SYSTEM_PROMPT),
        ChatMessage(
            role="user",
            content=(
                f"{context}{history_block}\n\n"
                f"Запрос пользователя:\n{user_question}\n\n"
                f"Сформируй SQL SELECT-запрос."
            ),
        ),
    ]
    print("\n" + "=" * 80)
    print("🧠 PROMPT ДЛЯ SQL-МОДЕЛИ")
    print("=" * 80)
    print("SYSTEM PROMPT:")
    print(STATISTICS_SQL_SYSTEM_PROMPT)
    print("-" * 80)
    print("USER MESSAGE:")
    print(
        f"{context}{history_block}\n\n"
        f"Запрос пользователя:\n{user_question}\n\n"
        f"Сформируй SQL SELECT-запрос."
    )
    print("=" * 80 + "\n")
    response = Settings.llm.chat(messages)
    return validate_statistics_sql(extract_sql_from_llm_response(response.message.content))


def execute_statistics_sql(sql: str) -> pd.DataFrame:
    conn = get_db_connection()
    try:
        return pd.read_sql_query(sql, conn)
    finally:
        conn.close()


def dataframe_to_llm_rows(df: pd.DataFrame, max_rows: int = 50) -> str:
    if df.empty:
        return "Результат пуст."

    lines = []
    for i, (_, row) in enumerate(df.head(max_rows).iterrows(), start=1):
        parts = []
        for col in df.columns:
            parts.append(f"{col}: {row[col]}")
        lines.append(f"Строка {i}: " + "; ".join(parts))

    return "\n".join(lines)


def generate_statistics_answer(user_question: str, sql: str, df: pd.DataFrame) -> str:
    rows_text = dataframe_to_llm_rows(df)
    messages = [
        ChatMessage(role="system", content=STATISTICS_ANSWER_SYSTEM_PROMPT),
        ChatMessage(
            role="user",
            content=(
                f"Описание структуры БД:\n{STATISTICS_SCHEMA_DESCRIPTION}\n\n"
                f"Запрос пользователя:\n{user_question}\n\n"
                f"SQL-запрос:\n{sql}\n\n"
                f"Результаты:\n{rows_text}"
            ),
        ),
    ]
    response = Settings.llm.chat(messages)
    return response.message.content


def generate_statistics_response(user_question: str, conversation_history: str = None) -> str:
    try:
        sql = generate_statistics_sql(user_question, conversation_history=conversation_history)
        print(f"📈 SQL statistics agent: {sql}")
        df = execute_statistics_sql(sql)
        print(f"📊 Statistics rows returned: {len(df)}")
        return generate_statistics_answer(user_question, sql, df)
    except Exception as e:
        print(f"Ошибка статистического агента: {e}")
        return f"Ошибка при обработке статистического запроса: {str(e)}"


# === Классификатор и RAG функции ===
def classify_query_tool(user_question: str, conversation_history: str = None) -> str:
    try:
        user_input = user_question
        if conversation_history:
            user_input = f"История диалога:\n{conversation_history}\n\nТекущий запрос:\n{user_question}"

        messages = [
            ChatMessage(role="system", content=CLASSIFIER_SYSTEM_PROMPT),
            ChatMessage(role="user", content=user_input)
        ]
        response = Settings.llm.chat(messages)
        query_type = response.message.content.strip().lower()
        return query_type if query_type in ["rag", "dialog", "statistics"] else "dialog"
    except Exception as e:
        print(f"Ошибка классификации: {e}")
        return "dialog"


def retrieve_rag_context(user_question: str) -> str:
    try:
        index = load_vector_index()
        if index is None:
            return "Ошибка: векторный индекс не загружен"

        retriever = index.as_retriever(similarity_top_k=4)
        nodes = retriever.retrieve(user_question)

        # === 🔍 DEBUG: Печать найденных узлов ===
        print(f"\n🔎 ЗАПРОС: {user_question}")
        print(f"📊 Найдено узлов: {len(nodes)}")
        for i, node in enumerate(nodes, 1):
            print(f"\n--- Узел #{i} (score: {node.score if hasattr(node, 'score') else 'N/A'}) ---")
            print(f"Текст для поиска: {node.get_content()[:300]}...")  # первые 300 символов
            print(f"Метаданные: { {k: v for k, v in node.metadata.items() if k.startswith('meta_')} }")
        print("-" * 80 + "\n")
        # === /DEBUG ===

        if not nodes:
            return "Не найдено релевантных документов."

        # Формируем контекст с добавлением метаданных
        context_parts = []
        meta_columns_display = {
            "meta_Наименование района": "Наименование района",
            "meta_Агроклиматические условия района": "Агроклиматические условия района",
            "meta_Ответственная организация": "Ответственная организация",
            "meta_Источник": "Источник"
        }

        for node in nodes:
            doc_text = node.get_content()
            meta_parts = []
            for meta_key, display_name in meta_columns_display.items():
                if meta_key in node.metadata and node.metadata[meta_key]:
                    meta_parts.append(f"{display_name}: {node.metadata[meta_key]}")

            if meta_parts:
                full_doc = f"{doc_text}\n" + "\n".join(meta_parts)
            else:
                full_doc = doc_text
            context_parts.append(full_doc)

        context = "\n\n---\n\n".join(context_parts)
        return context

    except Exception as e:
        return f"Ошибка при извлечении контекста: {str(e)}"


def generate_rag_response(user_question: str, context: str, conversation_history: str = None) -> str:
    try:
        history_instruction = ""
        if conversation_history:
            history_instruction = f"\n\nИстория диалога:\n{conversation_history}\n\nУчитывай историю выше при формировании ответа."

        full_system_prompt = RAG_SYSTEM_PROMPT + history_instruction + f"\n\nКонтекст из базы знаний:\n{context}"

        print(f"Контекст: {context}")
        # === /DEBUG ===

        messages = [
            ChatMessage(role="system", content=full_system_prompt),
            ChatMessage(role="user", content="Пользовательский запрос: " + user_question)
        ]
        response = Settings.llm.chat(messages)
        return response.message.content
    except Exception as e:
        return f"Ошибка генерации ответа: {str(e)}"


def process_query_simple(user_question: str, context_history: str = None) -> str:
    if not user_question.strip():
        return "Ошибка: пожалуйста, введите ваш запрос."

    print(f"\nОбработка запроса: '{user_question[:100]}{'...' if len(user_question) > 100 else ''}'")

    query_type = classify_query_tool(user_question, conversation_history=context_history)
    print(f"Тип запроса: {query_type}")

    if query_type == "rag":
        print("Запуск RAG-поиска...")
        rag_context = retrieve_rag_context(user_question)
        return generate_rag_response(user_question, rag_context, conversation_history=context_history)
    if query_type == "statistics":
        print("Запуск статистического агента...")
        return generate_statistics_response(user_question, conversation_history=context_history)

    print("Запуск диалогового режима...")
    return generate_dialog_response(user_question, conversation_history=context_history)


# === Утилиты для экспорта ===
def parse_html_table(html: str):
    """Преобразует HTML-таблицу в список списков + извлекает источники БЕЗ дублирования"""
    soup = BeautifulSoup(html, 'html.parser')
    tables = soup.find_all('table')

    result = []
    for table in tables:
        table_data = []
        # Заголовки
        header_row = table.find('thead')
        if header_row:
            headers = [th.get_text(strip=True) for th in header_row.find_all('th')]
            if headers:
                table_data.append(headers)
        # Тело таблицы
        tbody = table.find('tbody') or table
        for row in tbody.find_all('tr'):
            cells = [td.get_text(strip=True) for td in row.find_all(['td', 'th'])]
            if cells and any(cells):
                table_data.append(cells)
        if table:
            result.append(table_data)

    # === ИЗВЛЕЧЕНИЕ ИСТОЧНИКОВ: только один раз, без дублирования ===
    sources_text = ""

    # Ищем контейнер с ответом
    markdown_div = soup.find(class_='markdown-content') or soup

    # Находим последнюю таблицу
    all_tables = markdown_div.find_all('table')
    if not all_tables:
        return result, ""

    last_table = all_tables[-1]

    # Собираем ВСЁ текстовое содержимое ПОСЛЕ последней таблицы
    after_parts = []
    found_last_table = False

    # Проходим по всем элементам в порядке DOM
    for elem in markdown_div.find_all(string=True):  # текстовые узлы
        parent = elem.parent
        # Проверяем, является ли родитель частью последней таблицы
        if parent == last_table or parent.find_parent('table') == last_table:
            found_last_table = True
            continue
        # Если уже прошли таблицу — собираем текст
        if found_last_table:
            text = elem.strip()
            if text and text not in after_parts:  # избегаем дубликатов
                after_parts.append(text)

    if after_parts:
        # Собираем в строку, фильтруя пустые и дубли
        seen = set()
        unique_lines = []
        for line in after_parts:
            line_clean = line.strip()
            if line_clean and line_clean not in seen:
                seen.add(line_clean)
                unique_lines.append(line_clean)
        sources_text = '\n'.join(unique_lines)

    return result, sources_text.strip()


def _add_sources_to_docx(doc, sources_text: str):
    """Добавляет источники в DOCX как обычный текст (без маркеров списка), URL подчёркнут"""
    if not sources_text or not sources_text.strip():
        return

    # Пустая строка перед блоком источников
    doc.add_paragraph()

    lines = sources_text.strip().split('\n')
    header_added = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # === ОБРАБОТКА ЗАГОЛОВКА "Опорные источники" ===
        if 'опорные источники' in line.lower():
            if not header_added:
                doc.add_heading('Опорные источники:', level=3)
                header_added = True
            # Извлекаем остаток строки после заголовка
            rest = re.sub(r'опорные источники[:\s]*', '', line, flags=re.IGNORECASE).strip()
            if not rest:
                continue  # только заголовок, переходим к следующей строке
            line = rest  # обрабатываем остаток как обычный текст

        # === ОБЫЧНЫЙ ПАРАГРАФ (без маркера списка!) ===
        p = doc.add_paragraph()  # <--- УБРАНО: style='List Bullet'

        # === ОБРАБОТКА URL: делаем подчёркнутым, НЕ дублируя ===
        url_match = re.search(r'(https?://[^\s\]\[•\n]+)', line)

        if url_match:
            url = url_match.group(1)
            parts = line.split(url)
            # Текст до URL
            if parts[0].strip():
                p.add_run(parts[0].strip())
            # URL подчёркнутый
            url_run = p.add_run(url)
            url_run.font.underline = True
            # Текст после URL (если есть)
            if len(parts) > 1 and parts[1].strip():
                p.add_run(parts[1].strip())
        else:
            # Нет URL — добавляем весь текст как есть
            p.add_run(line)


def _add_sources_to_excel(worksheet, sources_text: str, start_row: int):
    """Добавляет источники в Excel: одна строка = одна ячейка, без дублирования"""
    if not sources_text or not sources_text.strip():
        return start_row

    lines = sources_text.strip().split('\n')
    header_added = False

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Заголовок — только один раз
        if 'опорные источники' in line.lower() and not header_added:
            worksheet.cell(row=start_row, column=1, value="Опорные источники:")
            worksheet.cell(row=start_row, column=1).font = openpyxl.styles.Font(bold=True)
            start_row += 1
            header_added = True
            # Если в строке есть ещё текст после заголовка — обрабатываем
            rest = re.sub(r'опорные источники[:\s]*', '', line, flags=re.IGNORECASE).strip()
            if rest:
                line = rest
            else:
                continue

        # Добавляем строку источника
        cell = worksheet.cell(row=start_row, column=1, value=line)
        start_row += 1

        # Если в строке есть URL — добавляем гиперссылку
        url_match = re.search(r'(https?://[^\s]+)', line)
        if url_match:
            try:
                cell.hyperlink = url_match.group(1)
                cell.style = "Hyperlink"
            except Exception:
                pass  # openpyxl может не поддерживать все гиперссылки

    return start_row


# === FastAPI ===
app = FastAPI(title="Climate Adaptation API")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:8000", "http://127.0.0.1:8000"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


class QuestionRequest(BaseModel):
    question: str
    context: Optional[str] = None  # История диалога от PHP
    conversation_id: Optional[int] = None  # ID диалога (для логики)


class QuestionResponse(BaseModel):
    answer: str
    status: str


class ApprovedMeasure(BaseModel):
    name: str
    mitigation: str | None = None
    adaptation: str
    relevance: str
    responsible: str
    source_question: str | None = None


class ExportRequest(BaseModel):
    content: str
    filename: str


class GenerateSQLRequest(BaseModel):
    prompt: str
    table_name: str = "climate_cases"


class StructuredDataRequest(BaseModel):
    prompt: str


@app.get("/")
async def root():
    return {"message": "Climate Adaptation API is running"}


@app.post("/ask", response_model=QuestionResponse)
async def ask_question_simple(request: QuestionRequest):
    try:
        # Передаем контекст истории в обработку
        answer = process_query_simple(request.question, context_history=request.context)
        return QuestionResponse(answer=answer, status="success")
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/health")
async def health_check():
    return {"status": "healthy"}


# === Одобрение + авто-ребилд ===
@app.post("/approve-measure")
async def approve_measure(measure: ApprovedMeasure):
    global APPROVED_MEASURES
    APPROVED_MEASURES.append(measure.dict())
    print(f"✅ Одобрено мероприятие: {measure.name}")

    # Запускаем ребилд в фоне (не блокируя ответ)
    threading.Thread(target=background_rebuild_index, daemon=True).start()

    return {"success": True, "message": "Добавлено. Ребилд запущен в фоне."}


# === (Опционально) ручной ребилд ===
@app.post("/rebuild-index")
async def manual_rebuild():
    threading.Thread(target=background_rebuild_index, daemon=True).start()
    return {"status": "started", "message": "Ребилд запущен в фоне"}


# === Экспорт в DOCX ===
@app.post("/export/docx")
async def export_docx(request: ExportRequest):
    try:
        tables, sources_text = parse_html_table(request.content)
        if not tables:
            raise HTTPException(status_code=400, detail="No tables found in content")

        doc = DocxDocument()
        doc.add_heading('Экспорт адаптационных мероприятий', level=1)

        for idx, table_data in enumerate(tables, 1):
            if idx > 1:
                doc.add_page_break()

            if len(table_data) < 2:
                continue

            # Создаём таблицу
            rows, cols = len(table_data), len(table_data[0])
            word_table = doc.add_table(rows=rows, cols=cols)
            word_table.style = 'Table Grid'

            # Заполняем данными
            for i, row_data in enumerate(table_data):
                for j, cell_text in enumerate(row_data[:cols]):
                    cell = word_table.cell(i, j)
                    cell.text = cell_text
                    if i == 0:  # заголовки
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.bold = True
                                run.font.size = Pt(10)
                    else:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                run.font.size = Pt(9)

            # Добавляем источники после каждой таблицы (сохраняя исходное форматирование)
            _add_sources_to_docx(doc, sources_text)

        # Сохраняем в буфер
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = request.filename if request.filename.endswith('.docx') else f"{request.filename}.docx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"DOCX export error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Export error: {str(e)}")


# === Экспорт в Excel ===
@app.post("/export/excel")
async def export_excel(request: ExportRequest):
    try:
        tables, sources_text = parse_html_table(request.content)
        if not tables:
            raise HTTPException(status_code=400, detail="No tables found in content")

        buffer = BytesIO()
        sheet_created = False

        with pd.ExcelWriter(buffer, engine='openpyxl') as writer:
            for idx, table_data in enumerate(tables, 1):
                if len(table_data) < 2:
                    continue

                # === ИСПРАВЛЕНИЕ: согласуем количество колонок ===
                headers = table_data[0]
                num_cols = len(headers)

                # Обрезаем каждую строку данных до количества колонок в заголовке
                cleaned_rows = []
                for row in table_data[1:]:
                    if len(row) >= num_cols:
                        cleaned_rows.append(row[:num_cols])
                    elif len(row) > 0:
                        # Если строк меньше, дополняем пустыми
                        padded = row + [''] * (num_cols - len(row))
                        cleaned_rows.append(padded)

                if not cleaned_rows:
                    continue

                # Создаём DataFrame с согласованными колонками
                df = pd.DataFrame(cleaned_rows, columns=headers)

                sheet_name = f'Table_{idx}'[:31]  # Excel limit: 31 char
                df.to_excel(writer, sheet_name=sheet_name, index=False)
                sheet_created = True

                # Форматирование
                worksheet = writer.sheets[sheet_name]
                # Жирные заголовки
                for cell in worksheet[1]:
                    if cell.value:
                        cell.font = openpyxl.styles.Font(bold=True)

                # Добавляем источники под таблицей
                if sources_text:
                    start_row = len(df) + 3  # +1 header +1 empty +1 sources header
                    start_row = _add_sources_to_excel(worksheet, sources_text, start_row)

            # === ИСПРАВЛЕНИЕ: гарантируем видимость хотя бы одного листа ===
            if not sheet_created:
                worksheet = writer.book.create_sheet(title='No_Data')
                worksheet['A1'] = 'Нет данных для экспорта'
                worksheet.sheet_state = 'visible'
            else:
                # Делаем первый лист активным и видимым
                if writer.book.worksheets:
                    writer.book.active = 0
                    writer.book.worksheets[0].sheet_state = 'visible'

        buffer.seek(0)
        filename = request.filename if request.filename.endswith('.xlsx') else f"{request.filename}.xlsx"

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{filename}"}
        )

    except HTTPException:
        raise
    except Exception as e:
        print(f"Excel export error: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Export error: {str(e)}")


@app.post("/generate-structured-data")
async def generate_structured_data(request: StructuredDataRequest):
    try:
        system_prompt = """
Ты — эксперт по SQL и климатическим адаптационным мероприятиям.
Пользователь описывает климатический кейс в свободной форме.

ПРАВИЛА:
1. Если в описании НЕ указаны хотя бы один из двух ключевых элементов:
   - Предлагаемые мероприятия/меры
   - Географическая привязка (регион/район)
   То верни JSON: {"error": "Не указаны мероприятия или регион"}

2. Если элементы есть, преобразуй описание в JSON структуру для таблицы 'climate_cases'.

ВЕРНИ ТОЛЬКО JSON в таком формате:
{
    "problem": "описание проблемы",
    "measure_name": "мероприятия",
    "mitigation_effect": "митигационный эффект",
    "adaptation_effect": "адаптационный эффект", 
    "district_name": "район",
    "climate_conditions": "агроклиматические условия района по плану: Среднегодовая температура. Сезонные средние (температура - зима, весна, лето, осень). Количество безморозных дней/период без ночных заморозков. Годовое количество осадков. Средняя влажность. Дождливые сезоны (напр., максимум весной/осенью). Экстремальные явления: есть ли явления и если да, то какие. Почвенные характеристики. Высота над уровнем моря. Климатический пояс. Есть ли водоемы рядом и какие, и есть ли риск затопления.)",
    "responsible_org": "ответственная организация",
    "source_url": "URL источника"
}

Если каких-то данных нет в описании пользователя — сделай реалистичные предположения на основе региона.

ПРИМЕР JSON:
{
    "problem": "городской остров тепла, нехватка воды и высокое энергопотребление",
    "measure_name": "Комплекс «Открытые сады»: создание экологического городского кластера с пассивными зданиями, геотермальной системой, фотоэлектрической станцией, зелёными крышами",
    "mitigation_effect": "Снижение выбросов CO₂, уменьшение потребления энергии",
    "adaptation_effect": "Стабилизация температуры, удержание дождевой воды, снижение эффекта городского острова тепла",
    "district_name": "Брно (Южноморавский край, Чехия)",
    "climate_conditions": "Среднегодовая температура — 9,7 °C. Зима: –1 °C, весна: 10 °C, лето: 20 °C, осень: 11 °C. Безморозный период — 170 дней. Осадки — 520 мм/год, максимум летом. Влажность — 73%. Почвы — суглинистые. Высота — 246 м. Климат — умеренно континентальный. Рядом река Свратка, риск локальных подтоплений.",
    "responsible_org": "Чешский фонд экологического партнёрства, муниципалитет Брно",
    "source_url": "https://www.adapterraawards.cz/en/databaze/2019/areal-otevrena-zahrada-v-brne"
}
Верни только JSON без пояснений.
        """

        messages = [
            ChatMessage(role="system", content=system_prompt),
            ChatMessage(role="user", content=f"Описание кейса: {request.prompt}")
        ]

        response = Settings.llm.chat(messages)
        json_str = response.message.content.strip()

        # Очистка от возможных меток кода
        json_str = json_str.replace('```json', '').replace('```', '').strip()

        # Парсинг JSON
        data = json.loads(json_str)
        print(data)
        return {
            "success": True,
            "data": data,
            "sql": "",
            "message": "Данные успешно сгенерированы"
        }

    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": f"Ошибка генерации данных: {str(e)}"
        }


@app.post("/execute-sql")
async def execute_sql(request: dict):
    try:
        sql = request.get("sql", "").strip()
        if not sql:
            raise ValueError("SQL запрос не может быть пустым")

        if not sql.lower().startswith("insert into"):
            raise ValueError("Разрешены только INSERT запросы")

        conn = psycopg2.connect(**PSYCOPG_DB_PARAMS)
        cursor = conn.cursor()
        cursor.execute(sql)
        conn.commit()

        rows_affected = cursor.rowcount

        cursor.close()
        conn.close()

        return {
            "success": True,
            "message": f"Успешно добавлено {rows_affected} записей",
            "rows_affected": rows_affected
        }

    except Exception as e:
        return {
            "success": False,
            "error": str(e),
            "message": f"Ошибка выполнения SQL: {str(e)}"
        }


if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8001)
